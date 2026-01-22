"""
Document Loader for RAG
Load v√† chunk d·ªØ li·ªáu t·ª´ schedules, news, announcements v√† info.docx

@author TBU AI Team
"""
from typing import List, Dict, Generator, Optional
import json
from datetime import datetime
import logging
import sys
import os

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from rag_config import CHUNK_SIZE, CHUNK_OVERLAP, INFO_DOCX_PATH

logger = logging.getLogger(__name__)


def chunk_text(
    text: str, 
    chunk_size: int = CHUNK_SIZE, 
    overlap: int = CHUNK_OVERLAP
) -> List[str]:
    """
    Split text into overlapping chunks
    T√°ch vƒÉn b·∫£n th√†nh c√°c ƒëo·∫°n nh·ªè c√≥ ph·∫ßn ch·ªìng l·∫•p
    
    Args:
        text: Text to chunk
        chunk_size: Maximum characters per chunk
        overlap: Number of overlapping characters
        
    Returns:
        List of text chunks
    """
    if not text or len(text.strip()) == 0:
        return []
    
    text = text.strip()
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # T√¨m ƒëi·ªÉm ng·∫Øt h·ª£p l√Ω n·∫øu ch∆∞a ƒë·∫øn cu·ªëi
        if end < len(text):
            # ∆Øu ti√™n ng·∫Øt t·∫°i: d·∫•u xu·ªëng d√≤ng k√©p > xu·ªëng d√≤ng > d·∫•u ch·∫•m > d·∫•u ph·∫©y
            best_break = end
            
            for sep in ['\n\n', '\n', '. ', ', ', ' ']:
                # T√¨m v·ªã tr√≠ ng·∫Øt trong n·ª≠a sau c·ªßa chunk
                search_start = start + chunk_size // 2
                idx = text.rfind(sep, search_start, end)
                
                if idx != -1:
                    best_break = idx + len(sep)
                    break
            
            end = best_break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Di chuy·ªÉn start v·ªõi overlap
        start = end - overlap if end - overlap > start else end
    
    return chunks


def format_schedule_for_embedding(schedule: Dict) -> str:
    """
    Format schedule record for embedding
    Chuy·ªÉn ƒë·ªïi schedule th√†nh text d·ªÖ embed v√† search
    
    Args:
        schedule: Schedule dictionary from database
        
    Returns:
        Formatted text string
    """
    from datetime import date as date_type, time as time_type
    
    # Parse date
    date_value = schedule.get('date', '')
    date_str = ''
    day_of_week = schedule.get('dayOfWeek', schedule.get('day_of_week', ''))
    
    days_vi = ['Th·ª© Hai', 'Th·ª© Ba', 'Th·ª© T∆∞', 'Th·ª© NƒÉm', 'Th·ª© S√°u', 'Th·ª© B·∫£y', 'Ch·ªß Nh·∫≠t']
    
    if date_value:
        if isinstance(date_value, date_type) and not isinstance(date_value, datetime):
            # Handle datetime.date object from pyodbc
            date_str = date_value.strftime('%d/%m/%Y')
            if not day_of_week:
                day_of_week = days_vi[date_value.weekday()]
        elif isinstance(date_value, str):
            try:
                date_obj = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
                date_str = date_obj.strftime('%d/%m/%Y')
                # Get Vietnamese day of week
                if not day_of_week:
                    day_of_week = days_vi[date_obj.weekday()]
            except:
                date_str = str(date_value)[:10]
        elif isinstance(date_value, datetime):
            date_str = date_value.strftime('%d/%m/%Y')
            if not day_of_week:
                day_of_week = days_vi[date_value.weekday()]
    
    # Format time
    start_time = schedule.get('startTime', schedule.get('start_time', ''))
    end_time = schedule.get('endTime', schedule.get('end_time', ''))
    
    # Handle datetime/time objects for time
    if isinstance(start_time, datetime):
        start_time = start_time.strftime('%H:%M')
    elif isinstance(start_time, time_type):
        start_time = start_time.strftime('%H:%M')
    elif isinstance(start_time, str) and 'T' in start_time:
        start_time = start_time.split('T')[1][:5]
    
    if isinstance(end_time, datetime):
        end_time = end_time.strftime('%H:%M')
    elif isinstance(end_time, time_type):
        end_time = end_time.strftime('%H:%M')
    elif isinstance(end_time, str) and 'T' in end_time:
        end_time = end_time.split('T')[1][:5]
    
    # Build text parts
    parts = []
    
    # Title/Header
    if date_str:
        parts.append(f"L·ªãch c√¥ng t√°c ng√†y {date_str} ({day_of_week})")
    
    # Time
    if start_time and end_time:
        parts.append(f"Th·ªùi gian: {start_time} - {end_time}")
    elif start_time:
        parts.append(f"Th·ªùi gian b·∫Øt ƒë·∫ßu: {start_time}")
    
    # Content
    content = schedule.get('content', '')
    if content:
        parts.append(f"N·ªôi dung: {content}")
    
    # Location
    location = schedule.get('location', '')
    if location:
        parts.append(f"ƒê·ªãa ƒëi·ªÉm: {location}")
    
    # Leader
    leader = schedule.get('leader', '')
    if leader:
        parts.append(f"Ch·ªß tr√¨: {leader}")
    
    # Participants
    participants = schedule.get('participants', [])
    if isinstance(participants, str):
        try:
            participants = json.loads(participants)
        except:
            participants = [participants] if participants else []
    
    if participants and isinstance(participants, list) and len(participants) > 0:
        parts.append(f"Th√†nh ph·∫ßn tham d·ª±: {', '.join(participants)}")
    
    # Preparing unit
    preparing_unit = schedule.get('preparingUnit', schedule.get('preparing_unit', ''))
    if preparing_unit:
        parts.append(f"ƒê∆°n v·ªã chu·∫©n b·ªã: {preparing_unit}")
    
    # Cooperating units
    coop_units = schedule.get('cooperatingUnits', schedule.get('cooperating_units', []))
    if isinstance(coop_units, str):
        try:
            coop_units = json.loads(coop_units)
        except:
            coop_units = []
    
    if coop_units and isinstance(coop_units, list) and len(coop_units) > 0:
        parts.append(f"ƒê∆°n v·ªã ph·ªëi h·ª£p: {', '.join(coop_units)}")
    
    # Notes
    notes = schedule.get('notes', '')
    if notes:
        parts.append(f"Ghi ch√∫: {notes}")
    
    return '\n'.join(parts)


def format_news_for_embedding(news: Dict) -> str:
    """
    Format news record for embedding
    
    Args:
        news: News dictionary from database
        
    Returns:
        Formatted text string
    """
    parts = []
    
    title = news.get('title', '')
    if title:
        parts.append(f"Tin t·ª©c: {title}")
    
    summary = news.get('summary', '')
    if summary:
        parts.append(f"T√≥m t·∫Øt: {summary}")
    
    content = news.get('content', '')
    if content:
        # Remove HTML tags if any
        import re
        clean_content = re.sub(r'<[^>]+>', '', content)
        parts.append(f"N·ªôi dung: {clean_content}")
    
    category = news.get('category', '')
    if category:
        parts.append(f"Danh m·ª•c: {category}")
    
    published_at = news.get('publishedAt', news.get('published_at', ''))
    if published_at:
        if isinstance(published_at, datetime):
            published_at = published_at.strftime('%d/%m/%Y')
        parts.append(f"Ng√†y ƒëƒÉng: {published_at}")
    
    return '\n'.join(parts)


def format_announcement_for_embedding(announcement: Dict) -> str:
    """
    Format announcement record for embedding
    
    Args:
        announcement: Announcement dictionary from database
        
    Returns:
        Formatted text string
    """
    parts = []
    
    title = announcement.get('title', '')
    if title:
        parts.append(f"Th√¥ng b√°o: {title}")
    
    content = announcement.get('content', '')
    if content:
        # Remove HTML tags if any
        import re
        clean_content = re.sub(r'<[^>]+>', '', content)
        parts.append(f"N·ªôi dung: {clean_content}")
    
    priority = announcement.get('priority', '')
    if priority:
        priority_map = {
            'urgent': 'Kh·∫©n c·∫•p',
            'important': 'Quan tr·ªçng', 
            'normal': 'B√¨nh th∆∞·ªùng'
        }
        parts.append(f"M·ª©c ƒë·ªô: {priority_map.get(priority, priority)}")
    
    published_at = announcement.get('publishedAt', announcement.get('published_at', ''))
    if published_at:
        if isinstance(published_at, datetime):
            published_at = published_at.strftime('%d/%m/%Y')
        parts.append(f"Ng√†y ƒëƒÉng: {published_at}")
    
    return '\n'.join(parts)


def load_info_docx() -> List[Dict]:
    """
    Load and chunk info.docx file
    
    Returns:
        List of dicts with 'content' and 'metadata' keys
    """
    if not INFO_DOCX_PATH.exists():
        logger.warning(f"‚ö†Ô∏è info.docx not found at {INFO_DOCX_PATH}")
        return []
    
    try:
        from docx import Document
        
        doc = Document(INFO_DOCX_PATH)
        
        # Extract text from paragraphs
        full_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() 
                    for cell in row.cells 
                    if cell.text.strip()
                )
                if row_text:
                    full_text.append(row_text)
        
        combined_text = '\n'.join(full_text)
        
        # Chunk the text
        chunks = chunk_text(combined_text)
        
        logger.info(f"üìÑ Loaded {len(chunks)} chunks from info.docx")
        
        return [
            {
                'content': chunk,
                'metadata': {
                    'source': 'info.docx',
                    'chunk_index': i
                }
            }
            for i, chunk in enumerate(chunks)
        ]
        
    except ImportError:
        logger.error("‚ùå python-docx not installed. Run: pip install python-docx")
        return []
    except Exception as e:
        logger.error(f"‚ùå Failed to load info.docx: {e}")
        return []


def load_schedules_from_db(cursor) -> Generator[Dict, None, None]:
    """
    Load schedules from database via cursor
    
    Args:
        cursor: Database cursor
        
    Yields:
        Schedule dictionaries
    """
    cursor.execute("""
        SELECT 
            id, date, day_of_week, start_time, end_time, 
            content, location, leader, participants, 
            preparing_unit, cooperating_units, notes, status
        FROM schedules 
        WHERE status = 'approved'
        ORDER BY date DESC
    """)
    
    columns = [col[0] for col in cursor.description]
    
    for row in cursor.fetchall():
        yield dict(zip(columns, row))


def load_news_from_db(cursor) -> Generator[Dict, None, None]:
    """
    Load news from database via cursor
    
    Args:
        cursor: Database cursor
        
    Yields:
        News dictionaries
    """
    cursor.execute("""
        SELECT id, title, summary, content, category, published_at
        FROM news
        ORDER BY published_at DESC
    """)
    
    columns = [col[0] for col in cursor.description]
    
    for row in cursor.fetchall():
        yield dict(zip(columns, row))


def load_announcements_from_db(cursor) -> Generator[Dict, None, None]:
    """
    Load announcements from database via cursor
    
    Args:
        cursor: Database cursor
        
    Yields:
        Announcement dictionaries
    """
    cursor.execute("""
        SELECT id, title, content, priority, published_at
        FROM announcements
        WHERE expires_at IS NULL OR expires_at > GETDATE()
        ORDER BY published_at DESC
    """)
    
    columns = [col[0] for col in cursor.description]
    
    for row in cursor.fetchall():
        yield dict(zip(columns, row))


# Test function
def test_document_loader():
    """Test document loader functionality"""
    print("Testing Document Loader...")
    
    # Test chunking
    test_text = "A" * 1000
    chunks = chunk_text(test_text, chunk_size=300, overlap=50)
    print(f"Chunked 1000 chars into {len(chunks)} chunks")
    
    # Test schedule formatting
    test_schedule = {
        'date': '2026-01-22',
        'startTime': '08:00',
        'endTime': '11:00',
        'content': 'H·ªçp Ban Gi√°m hi·ªáu',
        'location': 'Ph√≤ng h·ªçp A1',
        'leader': 'Hi·ªáu tr∆∞·ªüng',
        'participants': '["PHT1", "PHT2", "Tr∆∞·ªüng ph√≤ng"]',
        'preparingUnit': 'Ph√≤ng H√†nh ch√≠nh'
    }
    
    formatted = format_schedule_for_embedding(test_schedule)
    print(f"\nFormatted schedule:\n{formatted}")
    
    # Test docx loading
    docs = load_info_docx()
    print(f"\nLoaded {len(docs)} chunks from info.docx")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    test_document_loader()
