"""
Script to create sample info.docx for TBU RAG Chatbot
Run this to generate the document
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_info_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('THÔNG TIN TRƯỜNG ĐẠI HỌC THÁI BÌNH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== GIỚI THIỆU CHUNG =====
    doc.add_heading('1. GIỚI THIỆU CHUNG', level=1)
    doc.add_paragraph("""
Trường Đại học Thái Bình (Thai Binh University - TBU) là cơ sở giáo dục đại học công lập trực thuộc Ủy ban Nhân dân tỉnh Thái Bình. Trường được thành lập nhằm đào tạo nguồn nhân lực chất lượng cao phục vụ sự nghiệp phát triển kinh tế - xã hội của tỉnh Thái Bình và khu vực đồng bằng Bắc Bộ.

Trường Đại học Thái Bình có sứ mệnh đào tạo, bồi dưỡng nguồn nhân lực có phẩm chất chính trị, đạo đức tốt, có kiến thức chuyên môn vững vàng, có kỹ năng thực hành và khả năng thích ứng với môi trường làm việc trong nước và quốc tế.
""")

    # ===== THÔNG TIN LIÊN HỆ =====
    doc.add_heading('2. THÔNG TIN LIÊN HỆ', level=1)
    doc.add_paragraph("""
- **Tên trường**: Trường Đại học Thái Bình (Thai Binh University)
- **Địa chỉ**: Phố Lý Thường Kiệt, Phường Quang Trung, Thành phố Thái Bình, Tỉnh Thái Bình
- **Điện thoại**: (0227) 3831.299 - 3831.399
- **Fax**: (0227) 3831.499
- **Email**: info@tbu.edu.vn
- **Website**: https://tbu.edu.vn
- **Fanpage Facebook**: https://facebook.com/tbu.edu.vn
""")

    # ===== CƠ CẤU TỔ CHỨC =====
    doc.add_heading('3. CƠ CẤU TỔ CHỨC', level=1)
    
    doc.add_heading('3.1. Ban Giám hiệu', level=2)
    doc.add_paragraph("""
- Hiệu trưởng: PGS.TS. Nguyễn Văn A
- Phó Hiệu trưởng phụ trách Đào tạo: TS. Trần Văn B
- Phó Hiệu trưởng phụ trách NCKH & HTQT: TS. Lê Thị C
- Phó Hiệu trưởng phụ trách CSVC: ThS. Phạm Văn D
""")

    doc.add_heading('3.2. Các Phòng ban', level=2)
    doc.add_paragraph("""
1. Phòng Tổ chức - Hành chính
2. Phòng Đào tạo
3. Phòng Công tác sinh viên
4. Phòng Khoa học - Công nghệ
5. Phòng Hợp tác Quốc tế
6. Phòng Tài chính - Kế toán
7. Phòng Quản trị - Thiết bị
8. Phòng Thanh tra - Pháp chế
9. Phòng Khảo thí & Đảm bảo chất lượng
10. Trung tâm Thư viện - Thông tin
11. Trung tâm Tin học - Ngoại ngữ
""")

    doc.add_heading('3.3. Các Khoa', level=2)
    doc.add_paragraph("""
1. Khoa Kinh tế
2. Khoa Kế toán - Tài chính
3. Khoa Quản trị Kinh doanh
4. Khoa Công nghệ Thông tin
5. Khoa Nông nghiệp
6. Khoa Sư phạm
7. Khoa Ngoại ngữ
8. Khoa Luật
9. Khoa Xây dựng
10. Khoa Điện - Điện tử
""")

    # ===== NGÀNH ĐÀO TẠO =====
    doc.add_heading('4. CÁC NGÀNH ĐÀO TẠO', level=1)
    
    doc.add_heading('4.1. Đào tạo Đại học', level=2)
    doc.add_paragraph("""
**Khối ngành Kinh tế - Quản lý:**
- Quản trị kinh doanh
- Kế toán
- Tài chính - Ngân hàng
- Kinh tế
- Marketing

**Khối ngành Công nghệ:**
- Công nghệ thông tin
- Kỹ thuật phần mềm
- Hệ thống thông tin quản lý
- Kỹ thuật điện
- Kỹ thuật điện tử

**Khối ngành Nông nghiệp:**
- Khoa học cây trồng
- Chăn nuôi
- Nuôi trồng thủy sản
- Công nghệ thực phẩm

**Khối ngành Sư phạm:**
- Sư phạm Toán học
- Sư phạm Văn học
- Sư phạm Tiếng Anh
- Giáo dục Tiểu học
- Giáo dục Mầm non

**Khối ngành Xây dựng - Kiến trúc:**
- Kỹ thuật xây dựng
- Kiến trúc
- Quản lý xây dựng
""")

    doc.add_heading('4.2. Đào tạo Sau đại học', level=2)
    doc.add_paragraph("""
- Thạc sĩ Quản trị kinh doanh
- Thạc sĩ Kế toán
- Thạc sĩ Khoa học máy tính
- Thạc sĩ Nông nghiệp
""")

    # ===== TUYỂN SINH =====
    doc.add_heading('5. THÔNG TIN TUYỂN SINH', level=1)
    doc.add_paragraph("""
**Thời gian tuyển sinh**: Hàng năm từ tháng 3 đến tháng 9

**Phương thức xét tuyển:**
1. Xét tuyển dựa trên điểm thi THPT Quốc gia
2. Xét tuyển học bạ THPT
3. Xét tuyển kết hợp

**Chỉ tiêu tuyển sinh năm 2026**: Khoảng 3.000 sinh viên

**Học phí**: 
- Các ngành Kinh tế, Quản lý: 15-18 triệu đồng/năm
- Các ngành Kỹ thuật: 18-22 triệu đồng/năm
- Các ngành Sư phạm: Miễn học phí theo chính sách

**Hỗ trợ sinh viên:**
- Ký túc xá với sức chứa 2.000 sinh viên
- Học bổng khuyến khích học tập (từ 1-5 triệu đồng/học kỳ)
- Hỗ trợ vay vốn sinh viên
- Giới thiệu việc làm sau tốt nghiệp

**Liên hệ tuyển sinh:**
- Phòng Đào tạo: (0227) 3831.299 - số máy lẻ 101
- Email: tuyensinh@tbu.edu.vn
- Hotline: 0912.345.678
""")

    # ===== CƠ SỞ VẬT CHẤT =====
    doc.add_heading('6. CƠ SỞ VẬT CHẤT', level=1)
    doc.add_paragraph("""
Trường Đại học Thái Bình có cơ sở vật chất hiện đại:

- **Diện tích khuôn viên**: 50 hecta
- **Giảng đường**: 100 phòng học với sức chứa từ 50-300 sinh viên
- **Phòng thí nghiệm**: 30 phòng thực hành, thí nghiệm chuyên ngành
- **Thư viện**: Thư viện điện tử với hơn 100.000 đầu sách và tài liệu số
- **Ký túc xá**: 10 tòa nhà với sức chứa 2.000 sinh viên
- **Khu thể thao**: Sân vận động, nhà thi đấu đa năng, bể bơi
- **Căn tin**: 3 căn tin phục vụ ăn uống
- **WiFi**: Phủ sóng toàn khuôn viên trường
""")

    # ===== GIỜ LÀM VIỆC =====
    doc.add_heading('7. GIỜ LÀM VIỆC', level=1)
    doc.add_paragraph("""
**Giờ hành chính:**
- Buổi sáng: 7h30 - 11h30
- Buổi chiều: 13h30 - 17h00
- Từ thứ Hai đến thứ Sáu (nghỉ thứ Bảy và Chủ Nhật)

**Thư viện:**
- Buổi sáng: 7h00 - 11h30
- Buổi chiều: 13h00 - 21h00
- Hoạt động cả thứ Bảy và Chủ Nhật

**Phòng máy tính:**
- 7h00 - 22h00 hàng ngày
""")

    # ===== CÂU HỎI THƯỜNG GẶP =====
    doc.add_heading('8. CÂU HỎI THƯỜNG GẶP (FAQ)', level=1)
    doc.add_paragraph("""
**Q: Trường có những hình thức đào tạo nào?**
A: Trường có các hình thức: Đào tạo chính quy, Đào tạo liên thông, Đào tạo văn bằng 2, Đào tạo từ xa.

**Q: Sinh viên có được ở ký túc xá không?**
A: Có, sinh viên năm nhất được ưu tiên ở ký túc xá. Chi phí khoảng 200.000 - 400.000 đồng/tháng.

**Q: Trường có chương trình học bổng không?**
A: Có nhiều loại học bổng: Học bổng khuyến khích học tập, Học bổng tài năng, Học bổng doanh nghiệp tài trợ.

**Q: Làm thế nào để liên hệ với phòng ban?**
A: Gọi tổng đài (0227) 3831.299 và chọn số máy lẻ của phòng ban cần liên hệ.

**Q: Thời gian nghỉ hè, nghỉ Tết như thế nào?**
A: Nghỉ hè từ tháng 6 đến tháng 8. Nghỉ Tết Nguyên đán khoảng 2-3 tuần theo lịch chung.

**Q: Trường có hỗ trợ việc làm cho sinh viên sau tốt nghiệp không?**
A: Có, Trung tâm Hỗ trợ sinh viên thường xuyên tổ chức ngày hội việc làm và kết nối doanh nghiệp.
""")

    # Save document
    data_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(data_dir, 'data', 'info.docx')
    
    # Ensure data directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    print(f"✅ Created info.docx at: {output_path}")
    return output_path

if __name__ == "__main__":
    create_info_docx()
